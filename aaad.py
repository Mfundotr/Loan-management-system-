from customtkinter import *
from tkinter import *
from tkcalendar import *
from tkinter import messagebox
from docx import Document
from datetime import datetime,date,timedelta
import datetime
from docx.shared import Cm
()

import sqlite3
from tkinter import ttk
import customtkinter
customtkinter.set_appearance_mode("light")
try:
 connection = sqlite3.connect('customer.db')
 cursor=connection.cursor()
 cursor.execute("""CREATE TABLE customers(
 first_name Text,
 last_name Text,
 surname Text,
 gender Text,
 age Integer,
 id_number Integer,
 phone_number  Integer,
 account_number Integer,
 card_number Integer,
 bank_branch Text,
 work_place Text,
 oocupation Text,
 netsalary Integer,
 credit_request Integer,
 with_interest Integer,
 off_interest Integer,
 interest Integer,
 payment_Status Integer,
 not_paid_weeks Integer,
 payday Text,
 intrimited_month Text,
 salary_day Text,
 conctoller Text,
 file_N Text,
 application_Day Text,
 Status Text)
 """)
 connection.commit()
 connection.close()
except:
 pass
 

root=CTk()
root.geometry("500x400")
root.title("MIDDLE BEAM")
style=ttk.Style()


def write(*arg):
    
    firstname_info=firstname.get()
    lastname_info=lastname.get()
    surname_info=surname_STR.get()
    gender1_info=  text=checkvar1.get()
    gender2_info= text=checkvar2.get()
    kind= str(gender1_info) +str(gender2_info)
    age_info=age1.get()
    id_info=id_STR.get()
    phonenumber_info=str(0) + str(phonenumber.get())
    accountnumber_info=account_number_int.get()
    cardnumber_info=cardnumber_STR.get()
    clicked_info=clicked.get()
    creditrequest_info=credit_request_int.get()
    payday_info=pay_d.get_date()
    
    
    salary_info=salary_d.get_date()
    interest_info=int(interest1.get())
    controller_info=controller_STR.get()
    work_pc_info=work_pc_STR.get()
    occupation_info=occupation_STR.get()
    netsalary_info=netsalary_STR.get()
    name_relat_info=name_relat_STR.get()
    spause_info=spause_STR.get()
    resental_info=resental_STR.get()
  
    #primer detail oparation
    interest_info=interest1.get()
    
    interest_inxls=str(interest_info) + str("%")
    
    total=int(interest_info)/1*int(creditrequest_info) /100 +int(creditrequest_info)
    payment_s= total
   
    total2=int(interest_info)/1*int(creditrequest_info) /100
    with_interest= total2 + int(creditrequest_info)
    file_fn=firstname_info
    file_ln=lastname_info
    file_sn=surname_info
    file_id=id_info
    file_ids=str(file_id[10:13])
    file_fns=str((file_fn[0:1])) 
    file_lns=str((file_ln[0:1])) 
    file_sns=str((file_sn[0:1])) 
    files=file_fns+file_lns+file_sns+file_ids
    date=str(datetime.date.today())
      
    not_paid=int(0)
    dd=len(payday_info)
    if dd==7:
        sl1=(payday_info[0:1])
        dtt="0"+sl1
        sl2=(payday_info[2:4])
        sl3=(payday_info[5:7])
        fdt=dtt+"/"+sl2+"/"+sl3
        

        stuff= [(firstname_info,lastname_info,surname_info,kind , age_info , id_info, phonenumber_info, accountnumber_info,cardnumber_info,clicked_info,work_pc_info,occupation_info,netsalary_info,creditrequest_info,total,total2,interest_inxls,payment_s,not_paid,fdt,fdt,salary_info, controller_info,files,date,"PENDING",)]
        add_many(stuff)
        document=Document()

        header_section=document.sections[0]
        header=header_section.header
        header_text=header.paragraphs[0]
        header_text.text=" CASH LOANS(PTY) LTD\nREG NO:####/####/##\nP.O.BOX ##,########\nSTAND NO:###### ####\n#############"

        document.add_heading('                                     CASH LOAN AGREEMENT',1)
        document.add_paragraph('CONTROLLER INFO: ').add_run(controller_info.capitalize()).underline=True
        document.add_paragraph('CLIENT INFORMATION')
        n= document.add_paragraph('NAME AND SURNAME: ')
        n.add_run( firstname_info.capitalize()).underline=True
        n.add_run(" " +lastname_info.capitalize()).underline=True
        n.add_run(" " +surname_info.capitalize()).underline=True
        document.add_paragraph('CONTACT NUMBER: ').add_run( phonenumber_info).underline=True
        document.add_paragraph('WORKPLACE AND CONTACT: ').add_run(work_pc_info.capitalize()).underline=True
        document.add_paragraph('OCCUPATION: ').add_run(occupation_info.capitalize()).underline=True
        document.add_paragraph('NET SALARY: ').add_run(netsalary_info).underline=True
        document.add_paragraph('SALARY DAY: ').add_run( salary_info).underline=True
        sp= document.add_paragraph('NAME AND SURNAME OF SPAUSE/RELATIVE: ' )
        sp.add_run(spause_info.capitalize()).underline=True
        document.add_paragraph('CONTACT NUMBER OF SPAUSE/RELATIVE:  ').add_run( name_relat_info.capitalize()).underline=True
        document.add_paragraph('RESIDENTAL ADDRESS(school,clinic,shop,tavern,ect) ').add_run(resental_info.capitalize()).underline=True
        p=document.add_paragraph('FOR I ')
        p.add_run( firstname_info.capitalize()).underline=True
        p.add_run(" "+lastname_info.capitalize()).underline=True
        p.add_run(" "+surname_info.capitalize()).underline=True
        p.add_run(" ID ")
        p.add_run(str(" "+ id_info)).underline=True
        p.add_run(' HAVE BORROWERD FROM PARIS CASH LOAN ;ON THE ')
        p.add_run(date)
        p.add_run('  WITH AN INTEREST RATE OF ')
        p.add_run(str(interest_inxls))
        p.add_run(' TO BE CHARGED ON THE CAPITAL AMOUNT AND LEAVE MY ID BOOK')
        records={
            (creditrequest_info,total,fdt)
            }
        table = document.add_table(rows=1,cols=3)
        hdr_cells= table.rows[0].cells
        hdr_cells[0].text="AMOUNT BORROWED"
        hdr_cells[1].text="WITH_INTEREST"
        hdr_cells[2].text="PAY DAY" 
        for amount_borrowed,with_interest,fdt in records:
            row_cells=table.add_row().cells
            row_cells[0].text=str(amount_borrowed)
            row_cells[1].text=str(with_interest)
            row_cells[2].text=fdt
        document.add_paragraph('\nI DECLARE THAT I FAIL TO PAY THE BALANCE/AMOUNT DUE TO 7 DAYS AFTER THE DUE DATE,50% INTEREST RATE WOULD BE CHARGED ON THE AMONT DUE TO PROPERTY MAY BE TAKEN TO REPLACE THE BALANCE/AMONT DUE.')
        d= document.add_paragraph('SIGNED ON THE ')
        d.add_run(date).underline=True
        s=document.add_paragraph('CLIENT SIGNATURE_____________________')
        s.add_run('                 CONTROLLER SIGNATURE__________________')
    
        save_as=str(files) + str(".docx")
        document.save("C:\\Users\\Windows 10\Desktop\\loan application complete project\\Documents\\"+(save_as))
        #deleting all infor from widgts after fuction oparation complete
        firstname_entry.delete(0, END)
        lastname_entry.delete(0, END)
        surname_entry.delete(0, END)
        account_entry.delete(0 ,END)
        phonenumber_entry.delete(0, END)
        age_entry.delete(0,END) 
        credit_entry.delete(0, END)
        id_entry.delete(0,END) 
        interest_entry.delete(0,END)
        cardnumber_entry.delete(0,END)
        controller_entry.delete(0,END)
        work_pc_entry.delete(0,END)
        occpation_entry.delete(0,END)
        netsalary_entry.delete(0,END)
        spause_entry.delete(0,END)
        texxt.delete(0,END)
        name_relat_entry.delete(0,END)
        messagebox.showinfo("Done!","Application is sucessful")

    elif dd==6:
        sll1=(payday_info[0:1])
        dttt="0"+sll1
        slll2=(payday_info[2:3])
        dty="0"+slll2
        slll3=(payday_info[4:6])
        fdt=dttt+"/"+dty+"/"+slll3
        stuff= [(firstname_info,lastname_info,surname_info,kind ,age_info ,id_info, phonenumber_info, accountnumber_info,cardnumber_info,clicked_info,work_pc_info,occupation_info,netsalary_info,creditrequest_info,total,total2,interest_inxls,payment_s,not_paid,fdt,fdt,salary_info, controller_info,files,date,"PENDING",)]
        add_many(stuff)
        document=Document()

        header_section=document.sections[0]
        header=header_section.header
        header_text=header.paragraphs[0]
        header_text.text=" CASH LOANS(PTY) LTD\nREG NO:####/####/##\nP.O.BOX ##,########\nSTAND NO:###### ####\n#############"

        document.add_heading('                                     CASH LOAN AGREEMENT',1)
        document.add_paragraph('CONTROLLER INFO: ').add_run(controller_info.capitalize()).underline=True
        document.add_paragraph('CLIENT INFORMATION')
        n= document.add_paragraph('NAME AND SURNAME: ')
        n.add_run( firstname_info.capitalize()).underline=True
        n.add_run(" " +lastname_info.capitalize()).underline=True
        n.add_run(" " +surname_info.capitalize()).underline=True
        document.add_paragraph('CONTACT NUMBER: ').add_run( phonenumber_info).underline=True
        document.add_paragraph('WORKPLACE AND CONTACT: ').add_run(work_pc_info.capitalize()).underline=True
        document.add_paragraph('OCCUPATION: ').add_run(occupation_info.capitalize()).underline=True
        document.add_paragraph('NET SALARY: ').add_run(netsalary_info).underline=True
        document.add_paragraph('SALARY DAY: ').add_run( salary_info).underline=True
        sp= document.add_paragraph('NAME AND SURNAME OF SPAUSE/RELATIVE: ' )
        sp.add_run(spause_info.capitalize()).underline=True
        document.add_paragraph('CONTACT NUMBER OF SPAUSE/RELATIVE:  ').add_run( name_relat_info.capitalize()).underline=True
        document.add_paragraph('RESIDENTAL ADDRESS(school,clinic,shop,tavern,ect) ').add_run(resental_info.capitalize()).underline=True
        p=document.add_paragraph('FOR I ')
        p.add_run( firstname_info.capitalize()).underline=True
        p.add_run(" "+lastname_info.capitalize()).underline=True
        p.add_run(" "+surname_info.capitalize()).underline=True
        p.add_run(" ID ")
        p.add_run(str(" "+ id_info)).underline=True
        p.add_run(' HAVE BORROWERD FROM PARIS CASH LOAN ;ON THE ')
        p.add_run(date)
        p.add_run('  WITH AN INTEREST RATE OF ')
        p.add_run(str(interest_inxls))
        p.add_run(' TO BE CHARGED ON THE CAPITAL AMOUNT AND LEAVE MY ID BOOK')
        records={
            (creditrequest_info,total,fdt)
            }
        table = document.add_table(rows=1,cols=3)
        hdr_cells= table.rows[0].cells
        hdr_cells[0].text="AMOUNT BORROWED"
        hdr_cells[1].text="WITH_INTEREST"
        hdr_cells[2].text="PAY DAY"
        for amount_borrowed,with_interest,fdt in records:
            row_cells=table.add_row().cells
            row_cells[0].text=str(amount_borrowed)
            row_cells[1].text=str(with_interest)
            row_cells[2].text=fdt
        document.add_paragraph('\nI DECLARE THAT I FAIL TO PAY THE BALANCE/AMOUNT DUE TO 7 DAYS AFTER THE DUE DATE,50% INTEREST RATE WOULD BE CHARGED ON THE AMONT DUE TO PROPERTY MAY BE TAKEN TO REPLACE THE BALANCE/AMONT DUE.')
        d= document.add_paragraph('SIGNED ON THE ')
        d.add_run(date).underline=True
        s=document.add_paragraph('CLIENT SIGNATURE___________________')
        s.add_run('                 Controller SIGNATURE________________')
    
        save_as=str(files) + str(".docx")
        document.save("C:\\Users\\Windows 10\\Desktop\\loan application complete project\\Documents\\"+(save_as))
        #deleting all infor from widgts after fuction oparation complete
        firstname_entry.delete(0, END)
        lastname_entry.delete(0, END)
        surname_entry.delete(0, END)
        account_entry.delete(0 ,END)
        phonenumber_entry.delete(0, END)
        age_entry.delete(0,END) 
        credit_entry.delete(0, END)
        id_entry.delete(0,END) 
        interest_entry.delete(0,END)
        cardnumber_entry.delete(0,END)
        controller_entry.delete(0,END)
        work_pc_entry.delete(0,END)
        occpation_entry.delete(0,END)
        netsalary_entry.delete(0,END)
        spause_entry.delete(0,END)
        texxt.delete(0,END)
        name_relat_entry.delete(0,END)
        messagebox.showinfo("Done!","Application is sucessful")
                         
def home_page():
    home_frame= CTkFrame(main_frame,fg_color="transparent")
    global firstname
    global lastname
    global surname_STR
    global age
    global phonenumber
    global account_number_int
    global cardnumber_STR
    global credit_request_int
    global checkvar1
    global checkvar2 
    global age1
    global interest1
    global controller_STR
    global work_pc_STR
    global occupation_STR
    global netsalary_STR
    global name_relat_STR
    global name_relat_STR
    global resental_STR
    global clicked
    global resental_STR
    global spause_STR
    global id_STR
    global salary_d
    global pay_d
    
    firstname =StringVar(root)
    lastname=StringVar(root)
    surname_STR=StringVar()
    age=IntVar()
    phonenumber=StringVar()
    account_number_int=StringVar ()
    cardnumber_STR=StringVar()
    credit_request_int=StringVar()
    checkvar1 = StringVar()
    checkvar2 = StringVar()
    age1=StringVar()
    interest1=StringVar()
    controller_STR=StringVar()
    work_pc_STR=StringVar()
    occupation_STR=StringVar()
    netsalary_STR=StringVar()
    name_relat_STR=StringVar()
    name_relat_STR=StringVar()
    resental_STR=StringVar()
    clicked=StringVar()
    resental_STR=StringVar()
    spause_STR=StringVar()
    id_STR=StringVar()
    pd_d=StringVar()
    pd_m=StringVar()

    global firstname_entry
    global lastname_entry
    global surname_entry
    global age_entry
    global id_entry
    global phonenumber_entry
    global account_entry
    global cardnumber_entry
    global work_pc_entry
    global occpation_entry
    global netsalary_entry
    global name_relat_entry
    global spause_entry
    global texxt
    global controller_entry
    global credit_entry
    global interest_entry
 
    first_name= CTkLabel(home_frame,text="Enter First Name:")
    first_name.grid(row=15,column=1)
    firstname_entry=CTkEntry(home_frame,corner_radius=50,textvariable =firstname)
    firstname_entry.grid(row=20,column=1)

    last_name= CTkLabel (home_frame, text="Enter Last Name:")
    last_name.grid(row=25,column=1)
    
    lastname_entry=CTkEntry(home_frame,corner_radius=50,textvariable =lastname)
    lastname_entry.grid(row=30,column=1)
    surname =CTkLabel (home_frame, text ="Enter Surname:")
    surname.grid(row=35,column=1)
    
    surname_entry=CTkEntry(home_frame,corner_radius=50,textvariable =surname_STR )
    surname_entry.grid(row=40,column=1)
    
    c1= CTkCheckBox(home_frame,hover_color="green",fg_color="green",corner_radius=100,text = "Male", variable = checkvar1,onvalue = "Male" , offvalue = "" ,)
    c1.deselect()
    c1.grid(row=50,column=1)

       
    c2= CTkCheckBox(home_frame,hover_color="green",fg_color="green",corner_radius=100,text = "Female", variable = checkvar1,onvalue = "Female" , offvalue = "" ,)
    c2.deselect()
    c2.grid(row=50,column=2)


    age=CTkLabel(home_frame, text="Enter Age :")
    age.grid(row=55,column=1)
    age_entry=CTkEntry(home_frame,corner_radius=50,textvariable =age1)
    age_entry.grid(row=60,column=1)
    id=CTkLabel(home_frame, text="Enter ID Number:")
    id.grid(row=65,column=1)
    id_entry=CTkEntry(home_frame,corner_radius=50,textvariable=id_STR)
    id_entry.grid(row=70,column=1)
    phone_number=CTkLabel(home_frame, text="Phone Number :")
    phone_number.grid(row=75,column=1)
    phonenumber_entry=CTkEntry(home_frame,corner_radius=50, textvariable=phonenumber)
    phonenumber_entry.grid(row=80,column=1)

     
    account_number=CTkLabel(home_frame, text="Account Number :")
    account_number.grid(row=85,column=1)
    account_entry=CTkEntry(home_frame,corner_radius=50,textvariable=account_number_int)
    account_entry.grid(row=90,column=1)
    cardnumber=CTkLabel(home_frame,text="Card Number")
    cardnumber.grid(row=15,column=60)
    
    cardnumber_entry=CTkEntry(home_frame,corner_radius=50,textvariable=cardnumber_STR)
    cardnumber_entry.grid(row=20,column=60)
    work_pc=CTkLabel(home_frame,text="Workplace And Contact:")
    work_pc.grid(row=25,column=60)
    work_pc_entry= CTkEntry(home_frame,corner_radius=50,textvariable=work_pc_STR)
    work_pc_entry.grid(row=30,column=60)
    occupation=CTkLabel(home_frame,text="Occupation:")
    occupation.grid(row=35,column=60)
    occpation_entry=CTkEntry(home_frame,corner_radius=50,textvariable= occupation_STR)
    occpation_entry.grid(row=40,column=60)
    netsalary= CTkLabel(home_frame,text="Net Salary:")
    netsalary.grid(row=45,column=60)
    netsalary_entry= CTkEntry(home_frame,corner_radius=50,textvariable=netsalary_STR)
    netsalary_entry.grid(row=50,column=60)

    name_relative=CTkLabel(home_frame,text="Name of Spause/Relative")
    name_relative.grid(row=55,column=60)

    name_relat_entry=CTkEntry(home_frame,corner_radius=50,textvariable=name_relat_STR)
    name_relat_entry.grid(row=60,column=60)

    spause_l=CTkLabel(home_frame,text="Contact Number of Spause/Relative:")
    spause_l.grid(row=65,column=60)
    spause_entry = CTkEntry(home_frame,corner_radius=50,textvariable=spause_STR)
    spause_entry.grid(row=70,column=60)

    texxt= CTkEntry(home_frame,corner_radius=50,textvariable=resental_STR)
    texxt.grid(row=80,column=60)
    residental= CTkLabel(home_frame,text="Residental Address(School,Clinic,tavern ect:")
    residental.grid(row=75,column=60)
    dropl= CTkLabel(home_frame,text="Select Bank Branch:")
    dropl.grid(row=85,column=60)
   
    drop= CTkSegmentedButton(home_frame,corner_radius=25,width=25,selected_color="gray",selected_hover_color="gray",fg_color="green",unselected_color="green",unselected_hover_color="green",values=["African Bank","Capitec Bank","Absa Bank","Nedbank","Standard Bank","FNB Bank"],variable=clicked)
    drop.grid(row=90 ,column=60)
    
    controller_text= CTkLabel(home_frame,text="CONTROLER_INFO!")
    controller_text.grid(row=1,column=1)
    
    controller_entry= CTkEntry(home_frame,corner_radius=50,textvariable=controller_STR)
    controller_entry.grid(row=2,column=1)

    credit_request=CTkLabel(home_frame, text="Enter Credit_request:")
    credit_request.grid(row=30,column=78)

    credit_entry=CTkEntry (home_frame,corner_radius=50,textvariable =credit_request_int)
    credit_entry.grid(row=35,column=78)
    interests_f=CTkLabel(home_frame, text="Enter Interest:")
    interests_f.grid(row=40,column=78)
    interest_entry=CTkEntry(home_frame,corner_radius=50,textvariable=interest1)
    interest_entry.grid(row=45,column=78)

    salary=CTkLabel(home_frame,text="Salary Day")
    salary.grid(row=1,column=60)

    salary_d= Calendar(home_frame,selectmode="day", year=2024, month=3, day =1)
    salary_d.grid(row=2,column=60)

    
    pd=CTkLabel(home_frame,text="Pay Day")
    pd.grid(row=1,column=78)
    pay_d=Calendar(home_frame,selectmode="day", year=2024, month=3, day =1)
    pay_d.grid(row=2,column=78)
    
    button=CTkButton(home_frame,text ="Submit", command=write,corner_radius=50,fg_color="darkgreen",hover_color="green")
    button.grid(row=55,column=78)
    
    home_frame.pack()


         
def pay_page():
    home_frame=CTkFrame(main_frame,fg_color="transparent")

    index_idp=StringVar()
    amount_p=StringVar()
    
    
    label1=CTkLabel(home_frame, text="PAYING OPTION")
    label1.grid(row=0,column=1)
    label=CTkLabel(home_frame, text="ENTER ID NUMBER OF USER YOU WANT TO PAY TO")
    label.grid(row=2,column=1)
    index_number= CTkEntry(home_frame,corner_radius=50,textvariable =index_idp)
    index_number.grid(row=3,column=1)
    label1=CTkLabel(home_frame, text="ENTER AMOUNT DELLOW")
    label1.grid(row=4,column=1)
    pay_amount= CTkEntry(home_frame,corner_radius=50,textvariable =amount_p)
    pay_amount.grid(row=5,column=1)
  

    vv=CTkTextbox(home_frame,fg_color="transparent")
    vv.insert('0.0 ','PLSEASE USE CHECK BUTTOM TO CONFORM ID USER BEFORE TRANSATIONS') 
    vv.grid(row=1,column=1)
    def paying(*arg):
      connection = sqlite3.connect('customer.db')
      cursor=connection.cursor()
      date=str(datetime.date.today())
      try:
        money=int(amount_p.get())
        ffind=int(index_idp.get())
        if money ==' ':
          messagebox.showerror('Error','Please fill in ID field')
        if ffind ==' ':
            messagebox.showerror('Error','Please fill in ID field')
      except TclError:
         pass
      else:
              try:
                find= int(ffind)
              except ValueError:
                messagebox.showerror('Error','Use digits for ID prsenetion')
              try:
                n=messagebox.askquestion('Note','Do you want a to proceed')
                if n=='yes':
                  am=cursor.execute("SELECT payment_Status FROM customers WHERE id_number=?",[ffind])
                  uplode=cursor.execute("UPDATE customers SET Status='PEYING'  WHERE id_number=?",[ffind])
                  update=cursor.execute("UPDATE customers SET payment_Status=payment_status - ? WHERE id_number=?",(money,ffind))
                  fing=cursor.execute("SELECT payment_Status FROM customers WHERE id_number=?",[ffind])
                  for f in fing:
                    pass
                  
                elif n=='no':
                    pay_amount.delete(0,END)
                    index_number.delete(0,END)
                  
                  
                    
              except ValueError:
                pass
              try:
                
                messagebox.showinfo('Done','PAYMENT COMPLETE R'+ str(money)+ ' '+'IS PAID LEFT WITH R'+str( f)+ ' '+'TO PAY' )
              except UnboundLocalError:
                messagebox.showerror('Error','NO SUCH USER ID IN DATABASE')
      connection.commit()
      pay_amount.delete(0,END)
      index_number.delete(0,END)

    button1=CTkButton(home_frame,text ="PAY", command=paying,corner_radius=50,fg_color="darkgreen",hover_color="green")
    button1.grid(row=6,column=1)
    def chrk(*arg):
      connection = sqlite3.connect('customer.db')
      cursor=connection.cursor()
      date=str(datetime.date.today())
      try:
        ffind=index_idp.get()
        cgk1=cursor.execute("SELECT first_name FROM customers WHERE id_number=?",[ffind])
        for e in cgk1:
            for a in e:
                pass
        cgk2=cursor.execute("SELECT last_name FROM customers WHERE id_number=?",[ffind])
        for e in cgk2:
            for b in e:
                pass
        cgk3=cursor.execute("SELECT surname FROM customers WHERE id_number=?",[ffind])
        for e in cgk3:
            for c in e:
                pass
        cgk3=cursor.execute("SELECT gender FROM customers WHERE id_number=?",[ffind])
        for e in cgk3:
            for d in e:
                 
                
                vv.delete("0.0","end")
                z= "Names:  "+a+" "+b+" "+c+"\nGender: "+d
             
              
               
                vv.insert("0.0",z)
            
      except:
         messagebox.showerror('Error','NO SUCH USER ID IN DATABASE')
    button2=CTkButton(home_frame,text ="ID CHECK", command=chrk,corner_radius=50,fg_color="darkgreen",hover_color="green")
    button2.grid(row=6,column=3)

    
    home_frame.pack()



    
    
def check_page(*arg):
    home_frame=CTkFrame(main_frame,fg_color="transparent",height=500,width=800)
   
    home_framer=CTkScrollableFrame(home_frame,
    height=500,width=800,orientation="horizontal",
    label_text="USER INFO ",label_fg_color='green',
    label_text_color='white',
    scrollbar_button_color='green',
    scrollbar_button_hover_color='green',
    corner_radius=25
    )
    
    id_chk=StringVar()

    
    chk=CTkLabel(home_framer, text="ENTER ID NUMBER OR PHONE NUMBER")
    chk.grid(row=6,column=4)
    id_check= CTkEntry(home_framer,corner_radius=50,textvariable =id_chk)
    id_check.grid(row=7,column=4)
  


    tv= ttk.Treeview(home_framer,columns=(1,2,3,4,5),show="headings",height="4")
    

    tv.column('#0',anchor='center')
    
    tv.heading(1,text="FIRST NAME",anchor='w')
    tv.heading(2,text="LAST NAME",anchor='w')
    tv.heading(3,text="SURNAME",anchor='w')
    tv.heading(4,text="WITH int%",anchor='w')
    tv.heading(5,text="P-STATUS",anchor='w')
    tv.grid(row=20,column=5,rowspan=1,padx=5)
    def checking():
      connection = sqlite3.connect('customer.db')
      cursor=connection.cursor()
      id_checkn=float(id_chk.get())
      id_checktn=float(id_chk.get())
      checking_slq=cursor.execute("SELECT first_name,last_name,surname,with_interest,payment_status FROM customers WHERE id_number=? OR phone_number=?",(id_checkn,id_checktn))
      rows=cursor.fetchall()
      total=cursor.rowcount
      for i in rows:
        tv.insert('','end',values=i)
    button=CTkButton(home_framer,text ="CHECK", command=checking,corner_radius=50,fg_color="darkgreen",hover_color="green")
    button.grid(row=2,column=4
                )
    home_framer.grid(row=90,column=10)
    home_frame.pack()

   
  
    home_frame=CTkFrame(main_frame,fg_color="transparent")


    
    presentt=Frame(home_frame,height=160,width=160)


    style.map('Treeview',
              background=[('selected','green')])
    tgv= ttk.Treeview(presentt,columns=(1,2,3,4,5,6,7,8),show="headings",height="5")
    tgv.column('#0',anchor='center')
    tgv.column('#0',stretch=YES )
      
   
    def checkings():
      connection = sqlite3.connect('customer.db')
      cursor=connection.cursor()
      date=datetime.date.today()
      date2=(date.strftime("%m/%d/%y"))
      checking_slq=cursor.execute("SELECT first_name,last_name,surname,phone_number,gender,credit_request,with_interest,payment_status FROM customers WHERE payday=? AND payment_status>0",[date2])
      rows=cursor.fetchall()
      total=cursor.rowcount
      for i in rows:
        tgv.insert('','end',values=i)
      
    
      button=CTkButton(home_frame,text="SUBMIT", command=checkings,corner_radius=50,fg_color="green",hover_color="green")
      button.grid(row=5,column=10)



    presentt.grid(row=90,column=10)
    home_frame.pack()
   
def produce(*arg):
    connection = sqlite3.connect('customer.db')
    cursor=connection.cursor()
    date=str(datetime.date.today())
    try:
        ffind=int(index.get())
        
        if ffind ==' ':
            messagebox.showerror('Error','Please fill in ID field')
    except TclError:
            messagebox.showerror('Error','Use digits for ID prsenetion')
    else:
        try:
            ffind= int(ffind)

        except ValueError:
            messagebox.showerror('Error','Use digits for ID prsenetion')
    try:
        fn=cursor.execute("SELECT first_name FROM customers WHERE id_number=?",[ffind])
        data_row=fn.fetchall()
        for row in data_row:
            for ff in row:
                pass
        ln=cursor.execute("SELECT last_name FROM customers WHERE id_number=?",[ffind])
        data_row=ln.fetchall()
        for row in data_row:
            for l in row:
                pass
        ln=cursor.execute("SELECT surname FROM customers WHERE id_number=?",[ffind])
        data_row=ln.fetchall()
        for row in data_row:
            for s in row:
                pass
        ln=cursor.execute("SELECT id_number FROM customers WHERE id_number=?",[ffind])
        data_row=ln.fetchall()
        for row in data_row:
            for idd in row:
                pass
        ln=cursor.execute("SELECT with_interest FROM customers WHERE id_number=?",[ffind])
        data_row=ln.fetchall()
        for row in data_row:
            for w in row:
                pass
            
        ln=cursor.execute("SELECT payment_Status FROM customers WHERE id_number=?",[ffind])
        data_row=ln.fetchall()
        for row in data_row:
            for py in row:
                pass
        update=cursor.execute("UPDATE customers SET Status='Paid Full' WHERE id_number=?",[ffind])
        connection.commit()
        
    except UnboundLocalError:
        pass
    try:
        if py==0:
            document=Document()
            try:
                names=( ff + " " + l + " "+ s)
               
                header_section=document.sections[0]
                header=header_section.header
                header_text=header.paragraphs[0]
                header_text.text="PARIS CASH LOANS(PTY) LTD\nREG NO:2016/490927/07\nP.O.BOX 79,LUPHISA 1217\nSTAND NO:90LUPHISA 2017\n0799737227"
                hd=document.add_heading('\n                                     PAID UP LETTER',1)
                hd.underline=True
                document.add_paragraph('\nTO WHOM IT MAY CONSERN')
                document.add_paragraph('We confirm that the mentioned client has no outstanding balance with us Paris Cash Loans and that his/her loan has been paid in full')
                records={(names,idd,w,)
                         }
                table = document.add_table(rows=1,cols=3)
                hdr_cells= table.rows[0].cells
                hdr_cells[0].text="Client Info"
                hdr_cells[1].text="ID Number"
                hdr_cells[2].text="Amount "
                for names,iddd,credit_info in records:
                    row_cells=table.add_row().cells
                    row_cells[0].text=str(names)
                    row_cells[1].text=str(idd)
                    row_cells[2].text=str(w)
                regards= document.add_paragraph('Regards,')
                regards.underline=True
                document.add_paragraph('________________________')
                document.add_paragraph('branch manager')
                dt= document.add_paragraph(date)
                dt.underline=True
                document.add_page_break()
                file_fn=ff
                file_ln=l
                file_sn=s
                file_id=str(idd)
                file_ids=file_id[10:13]
                file_fns=str((file_fn[0:1]))
                file_lns=str((file_ln[0:1]))
                file_sns=str((file_sn[0:1]))
                files=file_fns+file_lns+file_sns+file_ids
                save_as=str(files) + str(".docx")
                document.save("C:\\Users\\Windows 10\\Desktop\\loan application complete project\\PFL Documents//" +(save_as))
                
                index_number.delete(0,END)
                messagebox.showinfo("DONE","Full Paid letter is produced")
               
            
            except UnboundLocalError:
                pass
        else:
            messagebox.showerror('Error','please make payment you have left with  R'+ str(py)+' '+ 'to pay')
    except:
        messagebox.showerror('Error','please type carefully')
        
               
def pfi_page():
    pfi_frame= CTkFrame(main_frame,height=100,width=100,fg_color="transparent" )
    global index
    global index_number
    index=StringVar()
    vv=CTkTextbox(pfi_frame,fg_color="transparent")
    vv.insert('0.0 ','THIS IS AN PAIDFULL PROGRAM USE ID NUMBER TO PRODUSE FINAL SETTLED LETTER') 
    vv.grid(row=1,column=1)
    label=CTkLabel(pfi_frame, text="Enter ID  number to produce letter ")
    label.grid(row=40,column=1)
    index_number= CTkEntry(pfi_frame,corner_radius=50,textvariable =index)
    index_number.grid(row=70,column=1)
    button=CTkButton(pfi_frame,text ="Submit", command=produce,corner_radius=50,fg_color="darkgreen",hover_color="green")
    button.grid(row=90,column=1)
    


    pfi_frame.pack()

         
def presento_page(*arg):
    home_frame=CTkFrame(main_frame,fg_color="transparent",height=500,width=1000)
    

    

    presento=CTkScrollableFrame(home_frame,height=500,width=800,orientation='horizontal',
    label_text="OUTSTANDING PROFILES OF TODAY ",label_fg_color='green',
    label_text_color='white',
    scrollbar_button_color='green',
    scrollbar_button_hover_color='green',
    corner_radius=25)

    style.map('Treeview',
            background=[('selected','green')])
    ti= ttk.Treeview(presento,columns=(1,2,3,4,5,6,7,8,9,10),show="headings",height="5")
    ti.column('#0',anchor='center')
    ti.column('#0',width=0,stretch=YES)
     
    ti.grid(row=15,column=0,rowspan=1,padx=5)
    ti.heading(1,text="FIRST(M)",anchor='w')
    ti.heading(2,text="LAST(N)",anchor='w')
    ti.heading(3,text="SURNAME",anchor='w')
    ti.heading(4,text="PHONE(N)",anchor='w')
    ti.heading(5,text="GENDER",anchor='w')
    ti.heading(6,text="CREDIT(R))",anchor='w')
    ti.heading(7,text="WITH(INT)",anchor='w')
    ti.heading(8,text="PAYMENT(S)",anchor='w')
    ti.heading(9,text="OWING(M)",anchor='w')
    ti.heading(10,text="NEW(PD)",anchor='w')
    
    def checkingo():
      connection = sqlite3.connect('customer.db')
      cursor=connection.cursor()
      date=datetime.date.today()
      duration2=timedelta(days=30)
      dff=date.strftime("%m/%d/%y")
      for d in range(30,duration2.days + 1):
        day2=date + timedelta(days=d)
        ddaye=(day2.strftime("%m/%d/%y"))
        print(ddaye)
        checking_slq=cursor.execute("SELECT first_name,last_name,surname,phone_number,gender,credit_request,with_interest,payment_status,not_paid_weeks,intrimited_month FROM customers WHERE intrimited_month =? OR intrimited_month=? AND payment_status>0 ",(ddaye,dff))
        rows=cursor.fetchall()
        for i in rows:
          ti.insert('','end',values=i)
    button=CTkButton(home_frame,text="SUBMIT", command=checkingo,corner_radius=50,fg_color="darkgreen",hover_color="green")
    button.grid(row=5,column=10)
    presento.grid(row=90,column=10)
    home_frame.pack()
    

def presentro_page(*arg):
    home_frame=CTkFrame(main_frame,fg_color="transparent",height=500,width=800)
   

    presentoo=CTkScrollableFrame(home_frame,
    height=500,width=800,orientation="horizontal",
    label_text="EVERY THING IN DATABASE ",label_fg_color='green',
    label_text_color='white',
    scrollbar_button_color='green',
    scrollbar_button_hover_color='green',
    corner_radius=25
    )
    
    
    
    style.map('Treeview',
            background=[('selected','green')])
    
    tri= ttk.Treeview(presentoo,columns=(1,2,3,4,5,6,7,8,9,10,11,12,13,14,15,16,17,18,19,20,21,22,23,24,25,26) ,show="headings",height="10")
   

   
 
    
    tri.column('#0',anchor='e')
    tri.column('#0',width=8,stretch=YES)
   
    tri.heading(1,text="FIRST NAME",anchor='w')
    tri.heading(2,text="LAST NAME",anchor='w')
    tri.heading(3,text="SURNAME",anchor='w')
    tri.heading(4,text="GENDER",anchor='w')
    tri.heading(5,text="AGE",anchor='w')
    tri.heading(6,text="ID",anchor='w')
    tri.heading(7,text="PHONE(N)",anchor='w')
    tri.heading(8,text="ACCOUNT(N) ",anchor='w')
    tri.heading(9,text="CARD(N) ",anchor='w')
    tri.heading(10,text="BANK  ",anchor='w')
    tri.heading(11,text="WORK(L)",anchor='w')
    tri.heading(12,text="OOCUPATION",anchor='w')
    tri.heading(13,text="NET SALARY",anchor='w')
    tri.heading(14,text="CREDIT(R)",anchor='w')
    tri.heading(15,text="WITH INTEREST ",anchor='w')
    tri.heading(16,text="Off INTEREST ",anchor='w')
    tri.heading(17,text="INTEREST",anchor='w')
    tri.heading(18,text="P_ STATUS",anchor='w')
    tri.heading(19,text="OWING MONTHS ",anchor='w')
    tri.heading(20,text="PAYDAY",anchor='w')
    tri.heading(21,text="DUE MONTH",anchor='w')
    tri.heading(22,text="SalaryDay",anchor='w')
    tri.heading(23,text="CONTROLLER",anchor='w')
    tri.heading(24,text="FILE(N)",anchor='w')
    tri.heading(25,text="APPLIED(M)",anchor='w')
    tri.heading(26,text="STATUS ",anchor='w')
    tri.grid(row=4,column=1,sticky='nsew')
   
    def go():
      connection = sqlite3.connect('customer.db')
      cursor=connection.cursor()
      checking_slq=cursor.execute("SELECT first_name,last_name,surname,gender,age,id_number,phone_number,account_number,card_number,bank_branch,work_place,oocupation,netsalary,credit_request,with_interest,off_interest,interest,payment_Status,not_paid_weeks ,payday,intrimited_month,salary_day,conctoller,file_N,application_Day,Status FROM customers")
      rows=cursor.fetchall()
      total=cursor.rowcount
      for i in rows:
          tri.insert('','end',values=i)
          connection.close()
          
   

    
    
    
    button=CTkButton(presentoo,text ="SUBMIT",command=go,corner_radius=50,fg_color="darkgreen",hover_color="green")
    button.grid(row=1,column=0)
    

    presentoo.grid(row=90,column=10)
    home_frame.pack()
    
    
connection = sqlite3.connect('customer.db')
cursor=connection.cursor()
    
date=datetime.date.today()
duration=timedelta(days=7)
for d in range(7,duration.days + 1):
    day=date - timedelta(days=d)
    ddye=(day.strftime("%m/%d/%y"))
    

    duration2=timedelta(days=30)
for d in range(30,duration2.days + 1):
     day2=date + timedelta(days=d)
     ddaye2=(day2.strftime("%m/%d/%y"))
     idp=cursor.execute("SELECT payment_Status FROM customers WHERE intrimited_month =?",[ddye])
     for d in idp:
         for r in d:
             rr= 50/100*r/1 + r
             
             upr=cursor.execute("UPDATE customers SET payment_Status= ?  WHERE intrimited_month =?",(rr,ddye))
           
             upr3=cursor.execute("UPDATE customers SET not_paid_weeks= not_paid_weeks+1  WHERE intrimited_month =?",[ddye])

        

             upr2=cursor.execute("UPDATE customers SET intrimited_month= ?  WHERE intrimited_month =?",(ddaye2,ddye))
             connection.commit()
             messagebox.showinfo("DONE","UNPAID CUSTOMERS GOT AN INCRESSE TODAY UPDATE ON TO PAY STATUS AMOUNT ")   

def delete_pages():
    for frame in main_frame.winfo_children():
        frame.destroy()
        
def indicate(lb,page):
   
    lb.configure()
    delete_pages()
    page()
options_frame=CTkFrame(root,height=500,fg_color="transparent")


home_btn=CTkButton(options_frame, text=" APPLY",corner_radius=50,fg_color="darkgreen",hover_color="green", font=("Bold",15),command=lambda:indicate(home_indicate,home_page))
home_btn.place(x=1,y=1)
home_indicate=CTkLabel(options_frame,text=' ')
home_indicate.place(x=3,y=50)

pay_btn=CTkButton(options_frame, text="PAY", corner_radius=50,fg_color="darkgreen",hover_color="green",font=("Bold",15),command=lambda:indicate(pay_indicate,pay_page))
pay_btn.place(x=10,y=50)
pay_indicate=CTkLabel(options_frame,text=' ')
pay_indicate.place(x=3,y=50)

check_btn=CTkButton(options_frame, text="CHECK",corner_radius=50,fg_color="darkgreen",hover_color="green", font=("Bold",15),command=lambda:indicate(check_indicate,check_page))
check_btn.place(x=10,y=100)
check_indicate=CTkLabel(options_frame,text=' ')
check_indicate.place(x=3,y=100)

paid_l_btn=CTkButton(options_frame, text="PFL",corner_radius=50,fg_color="darkgreen",hover_color="green", font=("Bold",15),command=lambda:indicate(paid_indicate,pfi_page))
paid_l_btn.place(x=6,y=150)
paid_indicate=CTkLabel(options_frame,text=' ', )
paid_indicate.place(x=3,y=150)

present_btn=CTkButton(options_frame, text="TODAY'S PROFILES",corner_radius=50,fg_color="darkgreen",hover_color="green", font=("Bold",15),command=lambda:indicate(paid_indicate,presento_page))
present_btn.place(x=6,y=200)
paid_indicate=CTkLabel(options_frame,text=' ', )
paid_indicate.place(x=3,y=200)

present_btn=CTkButton(options_frame, text="OUT STANDINGS",corner_radius=50,fg_color="darkgreen",hover_color="green", font=("Bold",15),command=lambda:indicate(paido_indicate,presento_page))
present_btn.place(x=6,y=250)

present_btn=CTkButton(options_frame, text="VIEW ALL",corner_radius=50,fg_color="darkgreen",hover_color="green", font=("Bold",15),command=lambda:indicate(paido_indicate,presentro_page))

present_btn.place(x=6,y=300)
paido_indicate=CTkLabel(options_frame,text=' ', )
paido_indicate.place(x=3,y=2500)

options_frame.pack(side=LEFT)
options_frame.pack_propagate(False)
options_frame.configure()

main_frame=CTkFrame(root,width=200,height=150,fg_color="transparent")
main_frame.pack(side=LEFT)
main_frame.pack(pady=10,padx=10,fill="both",expand=True)
main_frame.pack_propagate(False)
main_frame.configure()

connection = sqlite3.connect('customer.db')

cursor=connection.cursor()
date=datetime.date.today()
duration=timedelta(days=7)
for d in range(7,duration.days + 1):
    day=date - timedelta(days=d)
    ddye=(day.strftime("%m/%d/%y"))
    

duration2=timedelta(days=7)
for d in range(7,duration2.days + 1):
    day2=date + timedelta(days=d)
    ddaye2=(day2.strftime("%m/%d/%y"))
    

    idp=cursor.execute("SELECT payment_Status FROM customers WHERE intrimited_month =?",[ddye])
    for d in idp:
        for r in d:
            rr= 25/100*r/1 + r
            

            upr=cursor.execute("UPDATE customers SET payment_Status= ?  WHERE intrimited_month =?",(rr,ddye))
           
            upr3=cursor.execute("UPDATE customers SET not_paid_months= not_paid_months+1  WHERE intrimited_month =?",[ddye])

            upr2=cursor.execute("UPDATE customers SET intrimited_month= ?  WHERE intrimited_month =?",(ddaye2,ddye))
            connection.commit()
            messagebox.showinfo("DONE","UNPAID CUSTOMERS GOT AN INCRESSE TODAY UPDATE ON TO PAY STATUS AMOUNT ")
def add_many(list):
    connection =sqlite3.connect('customer.db')
    cursor= connection.cursor()
    cursor.executemany("INSERT INTO customers VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)",(list))
    connection.commit()
    connection.close()
root.mainloop()
